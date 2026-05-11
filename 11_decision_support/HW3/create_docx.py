import docx
from docx.shared import Inches
import re

def add_formatted_text(p, text):
    parts = re.split(r'(\*\*.*?\*\*|`.*?`|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Courier New'
        else:
            p.add_run(part)

def md_to_docx(md_filepath, docx_filepath):
    doc = docx.Document()
    
    with open(md_filepath, 'r') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        if line.startswith("# "):
            p = doc.add_heading(line[2:], 0)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], 1)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:], 2)
        elif line.startswith("* ") or line.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, line[2:])
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            # Remove "1. " etc.
            text = re.sub(r'^\d+\.\s', '', line)
            add_formatted_text(p, text)
        elif line.startswith("|") or line.startswith("---"):
            p = doc.add_paragraph()
            add_formatted_text(p, line)
        elif line.startswith("!["):
            # Image handling
            match = re.search(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                alt_text = match.group(1)
                img_path = match.group(2)
                try:
                    doc.add_picture(img_path, width=Inches(6.0))
                    doc.add_paragraph(f"Figure: {alt_text}")
                except Exception as e:
                    doc.add_paragraph(f"[Image placeholder for {img_path} - {e}]")
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line)
            
    doc.save(docx_filepath)
    print(f"Successfully saved {docx_filepath}")

if __name__ == "__main__":
    md_to_docx("HOMEWORK3_SUBMISSION.md", "HOMEWORK3_SUBMISSION.docx")
